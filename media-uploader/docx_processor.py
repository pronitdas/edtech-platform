import logging
import base64
import io
import re
import hashlib
import statistics
from typing import Dict, Tuple, List, Any, Optional, Set
from dataclasses import dataclass
from collections import defaultdict

import docx
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from PIL import Image

# Reuse TextBlock from pdf_processor.py
@dataclass
class TextBlock:
    """Represents a block of text with its properties."""
    text: str
    font_size: float
    font_name: str
    is_bold: bool
    is_italic: bool
    color: int
    bbox: Tuple[float, float, float, float]
    page_num: int
    block_type: str = "normal"  # Can be: title, heading, paragraph, list_item, table, etc.
    level: int = 0

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DOCXProcessor:
    """Processor for DOCX files with structure detection similar to PDFProcessor."""
    
    @staticmethod
    def extract_text_blocks(docx_document: DocxDocument) -> List[TextBlock]:
        """Extract text blocks with rich formatting information from the document."""
        all_blocks = []
        
        # Process paragraphs
        for page_num, paragraph in enumerate(docx_document.paragraphs):
            # Skip empty paragraphs
            if not paragraph.text.strip():
                continue
                
            # Process runs within the paragraph to get formatting
            for run in paragraph.runs:
                text = run.text.strip()
                if not text:
                    continue
                    
                # Extract formatting information
                font = run.font
                font_name = font.name if font.name else "Default"
                
                # Convert pt size to a comparable metric with PDF
                # In docx, size is in half-points, so divide by 2 to get points
                font_size = font.size.pt if font.size else 11.0  # Default size is typically 11pt
                
                is_bold = bool(font.bold)
                is_italic = bool(font.italic)
                
                # Simplified color handling (using a default color code)
                color = 0  # Default black
                if font.color and font.color.rgb:
                    # Convert RGB to an integer for consistency with PDF processor
                    color = int(font.color.rgb, 16)
                
                # Bounding box is not directly available in docx
                # Use default values for consistency with the interface
                bbox = (0, 0, 0, 0)
                
                # Create TextBlock
                all_blocks.append(TextBlock(
                    text=text,
                    font_size=font_size,
                    font_name=font_name,
                    is_bold=is_bold,
                    is_italic=is_italic,
                    color=color,
                    bbox=bbox,
                    page_num=page_num,
                    block_type="normal",
                    level=0
                ))
        
        return all_blocks
    
    @staticmethod
    def analyze_document_structure(blocks: List[TextBlock]) -> List[TextBlock]:
        """
        Analyze document structure based on styles and formatting.
        Reuses logic from PDFProcessor with adaptations for DOCX specifics.
        """
        if not blocks:
            return blocks
            
        # Analyze font size distribution
        font_sizes = [block.font_size for block in blocks if block.font_size > 0]
        if not font_sizes:
            return blocks
            
        try:
            # Find the most common font size (likely body text)
            body_font_size = statistics.mode(font_sizes)
        except statistics.StatisticsError:
            # If there's no clear mode, use median
            body_font_size = statistics.median(font_sizes)
        
        # Identify potential headings based on font properties
        size_threshold = body_font_size * 1.1  # 10% larger than body text
        
        # Collect all sizes larger than body text for heading level assignment
        heading_sizes = sorted(set(size for size in font_sizes if size > size_threshold), reverse=True)
        heading_levels = {size: idx + 1 for idx, size in enumerate(heading_sizes)}
        
        # Apply classification
        for block in blocks:
            # Rule 1: Size-based classification
            if block.font_size in heading_levels:
                block.block_type = "heading"
                block.level = heading_levels[block.font_size]
                continue
                
            # Rule 2: Bold text that's not common might be a heading
            bold_blocks = [b for b in blocks if b.is_bold]
            bold_percentage = len(bold_blocks) / len(blocks) if blocks else 0
            
            if block.is_bold and bold_percentage < 0.3:  # If less than 30% of text is bold
                if block.font_size >= body_font_size:
                    block.block_type = "heading"
                    block.level = len(heading_levels) + 1 if heading_levels else 1
                    continue
            
            # Rule 3: Content-based pattern matching
            text = block.text.strip()
            
            # Check for common heading patterns (reused from PDFProcessor)
            if re.match(r'^(?:chapter|section|part|appendix|figure|table)\s+\d+', text, re.IGNORECASE):
                block.block_type = "heading"
                block.level = 1 if text.lower().startswith(("chapter", "part")) else 2
                continue
                
            # Numbered headings
            if re.match(r'^\d+(\.\d+)*\.?\s+\S+', text):
                parts = re.match(r'^\d+(\.\d+)*', text).group(0).count('.') + 1
                block.block_type = "heading"
                block.level = parts
                continue
                
            # List items (reused from PDFProcessor)
            if re.match(r'^\s*[-•*]\s+\S+', text) or re.match(r'^\s*\d+\.\s+\S+', text):
                block.block_type = "list_item"
                continue
                
            # Default is paragraph
            block.block_type = "paragraph"
        
        # Title detection on first page
        if blocks:
            first_blocks = sorted([b for b in blocks if b.page_num == 0], 
                                key=lambda b: b.font_size, reverse=True)
            if first_blocks and first_blocks[0].font_size > body_font_size * 1.3:
                first_blocks[0].block_type = "title"
                
        return blocks
    
    @staticmethod
    def organize_blocks_into_document(blocks: List[TextBlock]) -> Dict:
        """
        Organize the classified blocks into a hierarchical document structure.
        This method is reused from PDFProcessor with no changes.
        """
        # Sort blocks by page number and vertical position
        sorted_blocks = sorted(blocks, key=lambda b: (b.page_num, b.bbox[1]))
        
        # Root structure
        document = {
            "title": "",
            "content": "",
            "children": []
        }
        
        # Find document title
        title_block = next((b for b in sorted_blocks if b.block_type == "title"), None)
        if title_block:
            document["title"] = title_block.text
        
        # Current section stack - [(level, node), ...]
        section_stack = [(0, document)]
        current_content = []
        
        # Process blocks
        for block in sorted_blocks:
            # Skip title block as it's handled separately
            if block.block_type == "title":
                continue
                
            if block.block_type == "heading":
                # Flush any accumulated content to current section
                if current_content:
                    section_stack[-1][1]["content"] += "\n".join(current_content)
                    current_content = []
                
                # Create new section
                new_section = {
                    "title": block.text,
                    "content": "",
                    "children": [],
                    "level": block.level,
                    "type": "section"
                }
                
                # Adjust stack - pop until we find a parent with a lower level
                while section_stack and section_stack[-1][0] >= block.level:
                    section_stack.pop()
                
                # Add new section to parent's children
                section_stack[-1][1]["children"].append(new_section)
                
                # Push new section onto stack
                section_stack.append((block.level, new_section))
            else:
                # Accumulate content
                current_content.append(block.text)
        
        # Flush any remaining content
        if current_content:
            section_stack[-1][1]["content"] += "\n".join(current_content)
        
        return document
    
    @staticmethod
    def extract_images(docx_document: DocxDocument) -> Dict[str, Any]:
        """Extract images from a DOCX document."""
        images = {}
        image_index = 0
        
        # Process document parts for images
        for rel in docx_document.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    # Get image data
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    
                    # Create a PIL image to get properties
                    image_stream = io.BytesIO(image_bytes)
                    pil_image = Image.open(image_stream)
                    
                    # Determine image format
                    image_format = pil_image.format.lower() if pil_image.format else "png"
                    
                    # Create unique image ID
                    image_hash = hashlib.md5(image_bytes).hexdigest()[:10]
                    image_filename = f"img_{image_index}_{image_hash}.{image_format}"
                    
                    # Check if image is too small (likely an icon or bullet)
                    if pil_image.width < 20 or pil_image.height < 20:
                        continue
                        
                    # Convert to base64
                    img_b64 = base64.b64encode(image_bytes).decode()
                    
                    # Store image with metadata
                    images[image_filename] = {
                        "data": img_b64,
                        "format": image_format,
                        "width": pil_image.width,
                        "height": pil_image.height,
                        "page": 0,  # DOCX doesn't have explicit pages like PDF
                        "mode": pil_image.mode,
                        "hash": image_hash
                    }
                    
                    image_index += 1
                except Exception as img_error:
                    logger.error(f"Error extracting image {image_index}: {str(img_error)}")
        
        return images
    
    @staticmethod
    def document_to_text(document: Dict) -> str:
        """
        Convert document structure to flat text with markdown formatting.
        This method is reused from PDFProcessor with no changes.
        """
        parts = []
        
        # Add title
        if document.get("title"):
            parts.append(f"# {document['title']}\n")
        
        # Add content
        if document.get("content"):
            parts.append(document["content"])
        
        # Process children recursively
        def process_section(section, depth=1):
            section_parts = []
            
            # Add heading with appropriate number of #
            heading_marks = "#" * (depth + 1)  # +1 because title is h1
            section_parts.append(f"\n{heading_marks} {section['title']}\n")
            
            # Add content
            if section.get("content"):
                section_parts.append(section["content"])
            
            # Process subsections
            for child in section.get("children", []):
                section_parts.append(process_section(child, depth + 1))
            
            return "\n".join(section_parts)
        
        # Process all top-level sections
        for section in document.get("children", []):
            parts.append(process_section(section))
        
        return "\n".join(parts)
    
    @staticmethod
    def process_docx(file_data: bytes) -> Tuple[str, Dict[str, Any], Dict[str, Any]]:
        """
        Process a DOCX file with structure detection.
        Returns: (text_content, images_dict, metadata_dict)
        """
        try:
            # Load the document from bytes
            docx_stream = io.BytesIO(file_data)
            docx_document = docx.Document(docx_stream)
            
            # Extract metadata
            metadata = {
                "title": docx_document.core_properties.title or "",
                "author": docx_document.core_properties.author or "",
                "subject": docx_document.core_properties.subject or "",
                "keywords": docx_document.core_properties.keywords or "",
                "creator": docx_document.core_properties.author or "",
                "format": "DOCX",
                "pages": len(docx_document.paragraphs),  # Approximation
                "file_size": len(file_data),
            }
            
            # Extract rich text blocks with formatting
            text_blocks = DOCXProcessor.extract_text_blocks(docx_document)
            
            # Analyze document structure
            classified_blocks = DOCXProcessor.analyze_document_structure(text_blocks)
            
            # Organize into hierarchical document
            document_structure = DOCXProcessor.organize_blocks_into_document(classified_blocks)
            
            # Extract images
            images = DOCXProcessor.extract_images(docx_document)
            
            # Generate flat text representation
            text_content = DOCXProcessor.document_to_text(document_structure)
            
            return text_content, images, metadata
            
        except Exception as e:
            logger.error(f"Error converting DOCX: {str(e)}")
            # Fallback to basic text extraction
            try:
                docx_stream = io.BytesIO(file_data)
                docx_document = docx.Document(docx_stream)
                text_content = "\n".join([para.text for para in docx_document.paragraphs])
                
                return text_content, {}, {"format": "DOCX"}
            except:
                raise  # Re-raise if even the fallback fails

    # Reuse methods from PDFProcessor for textbook conversion
    @staticmethod
    def parse_to_textbook(text: str) -> Dict:
        """Reuse the PDF textbook parsing logic for consistent output structure."""
        from pdf_processor import PDFProcessor
        return PDFProcessor.parse_pdf_to_textbook(text)
    
    @staticmethod
    def convert_tree_to_textbook(tree: Dict, knowledge_id: int, knowledge_name: str) -> Dict:
        """Reuse the PDF textbook conversion logic."""
        from pdf_processor import PDFProcessor
        return PDFProcessor.convert_tree_to_textbook(tree, knowledge_id, knowledge_name)
    
    @staticmethod
    def walk_textbook(textbook: Dict, knowledge_id: int) -> List[Dict]:
        """Reuse the PDF textbook walking logic."""
        from pdf_processor import PDFProcessor
        return PDFProcessor.walk_textbook(textbook, knowledge_id)
    
    @staticmethod
    def process_text_to_index(text: str, knowledge_id: int, knowledge_name: str) -> Tuple[Dict, List[Dict]]:
        """Reuse the PDF text to index processing logic."""
        from pdf_processor import PDFProcessor
        return PDFProcessor.process_pdf_text_to_index(text, knowledge_id, knowledge_name)
    
    @staticmethod
    def prepare_images_for_upload(images: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
        """
        Prepare extracted images for upload - custom implementation for DOCX.
        
        Args:
            images: Dictionary of images extracted from DOCX
            
        Returns:
            Dictionary ready for upload with standardized metadata and buffer
        """
        upload_ready = {}
        
        for filename, img_data in images.items():
            # Extract the base64 encoded data
            img_base64 = img_data.get("data", "")
            
            # Convert base64 to bytes for upload
            img_bytes = base64.b64decode(img_base64) if img_base64 else b""
            
            # Get image format
            img_format = img_data.get("format", "png")
            
            # Create standardized metadata
            upload_ready[filename] = {
                "buffer": img_bytes,  # Add the bytes data as buffer
                "format": img_format,
                "width": img_data.get("width", 0),
                "height": img_data.get("height", 0),
                "page": img_data.get("page", 1),
                "metadata": {
                    "filename": filename,
                    "format": img_format,
                    "width": img_data.get("width", 0),
                    "height": img_data.get("height", 0),
                    "page": img_data.get("page", 1),
                    "mode": img_data.get("mode", ""),
                    "hash": img_data.get("hash", ""),
                    "size_bytes": len(img_bytes) if img_bytes else 0
                }
            }
        
        return upload_ready 